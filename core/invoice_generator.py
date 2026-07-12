"""Honorarrechnungs-Generator (DOCX, RTF, TXT)."""

from datetime import datetime
from typing import Any

from core.config import get_settings
from core.models import FinalInvoicePosition, GeneratedInvoice, NotaryProfile


def generate_invoice(
    final_positions: list[dict[str, Any]],
    notary: dict[str, Any],
    original_document: str = "",
    output_format: str = "docx",
    fee_engine_version: str = "GNotKG_Stand_2026-01-01_v1",
) -> tuple[bytes, GeneratedInvoice]:
    """Erzeugt eine Honorarrechnung im gewählten Format. Gibt (bytes, GeneratedInvoice) zurück."""
    settings = get_settings()
    now = datetime.now()

    vat_rate = settings.app_vat_rate
    total_net = sum(p.get("fee_amount", 0.0) for p in final_positions)
    vat_amount = round(total_net * vat_rate, 2)
    total_gross = total_net + vat_amount

    if output_format == "docx":
        content = _generate_docx(
            final_positions,
            notary,
            now,
            total_net,
            vat_amount,
            total_gross,
            original_document,
            fee_engine_version,
        )
    elif output_format == "rtf":
        content = _generate_rtf(
            final_positions,
            notary,
            now,
            total_net,
            vat_amount,
            total_gross,
            original_document,
            fee_engine_version,
        )
    else:
        content = _generate_txt(
            final_positions,
            notary,
            now,
            total_net,
            vat_amount,
            total_gross,
            original_document,
            fee_engine_version,
        )

    positions = []
    for p in final_positions:
        positions.append(
            FinalInvoicePosition(
                kv_number=p.get("kv_number", ""),
                description=p.get("description", ""),
                business_value_eur=p.get("business_value_eur"),
                fee_amount=p.get("fee_amount", 0.0),
                source_reference=p.get("source_reference", ""),
                was_overridden=p.get("was_overridden", False),
                override_reason=p.get("override_reason"),
                calculation_details=p.get("reasoning", ""),
            )
        )

    profile = NotaryProfile(
        name=notary.get("name", ""),
        firm_name=notary.get("firm_name", ""),
        address=notary.get("address", ""),
        phone=notary.get("phone"),
        email=notary.get("email"),
        bank_name=notary.get("bank_name", ""),
        iban=notary.get("iban", ""),
        bic=notary.get("bic"),
        tax_number=notary.get("tax_number"),
        vat_id=notary.get("vat_id"),
    )

    invoice = GeneratedInvoice(
        created_at=now,
        notary=profile,
        original_document=original_document,
        positions=positions,
        total_net=total_net,
        vat_amount=vat_amount,
        total_gross=total_gross,
        output_formats=[output_format],  # type: ignore
        fee_engine_version=fee_engine_version,
    )
    return content, invoice


def _generate_docx(
    positions: list[dict],
    notary: dict,
    now: datetime,
    total_net: float,
    vat: float,
    total_gross: float,
    original_document: str,
    fee_version: str,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)

    # Notar-Kopf
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(notary.get("firm_name", "")).bold = True
    p.add_run("\n" + notary.get("name", ""))
    p.add_run("\n" + notary.get("address", ""))
    if notary.get("phone"):
        p.add_run("\nTel: " + notary["phone"])
    if notary.get("email"):
        p.add_run("\nE-Mail: " + notary["email"])

    doc.add_paragraph()

    # Titel
    title = doc.add_heading("Honorarrechnung", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rechnungsdatum
    doc.add_paragraph(f"Datum: {now.strftime('%d.%m.%Y')}")
    if original_document:
        doc.add_paragraph(f"Urkunde: {original_document}")

    doc.add_paragraph()

    # Positionen-Tabelle
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "KV-Nr."
    header[1].text = "Beschreibung"
    header[2].text = "Geschäftswert"
    header[3].text = "Gebühr (EUR)"
    header[4].text = "Fundstelle"

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for pos in positions:
        row = table.add_row()
        row.cells[0].text = pos.get("kv_number", "")
        row.cells[1].text = pos.get("description", "")
        bw = pos.get("business_value_eur")
        row.cells[2].text = f"{bw:,.2f} €" if bw else "–"
        row.cells[3].text = f"{pos.get('fee_amount', 0.0):,.2f} €"
        row.cells[4].text = pos.get("source_reference", "")

    doc.add_paragraph()

    # Summen
    summary = doc.add_paragraph()
    summary.add_run(f"Zwischensumme: {total_net:,.2f} €\n").bold = False
    summary.add_run(f"Umsatzsteuer 19 %: {vat:,.2f} €\n")
    run_total = summary.add_run(f"Gesamtbetrag: {total_gross:,.2f} €")
    run_total.bold = True

    doc.add_paragraph()

    # Zahlungsinformation
    payment = doc.add_paragraph()
    payment.add_run("Zahlung bitte auf folgendes Konto:\n").bold = True
    payment.add_run(f"{notary.get('bank_name', '')}\n")
    payment.add_run(f"IBAN: {notary.get('iban', '')}\n")
    if notary.get("bic"):
        payment.add_run(f"BIC: {notary.get('bic', '')}\n")
    if notary.get("tax_number"):
        payment.add_run(f"Steuernummer: {notary['tax_number']}\n")
    if notary.get("vat_id"):
        payment.add_run(f"USt-ID: {notary['vat_id']}\n")

    doc.add_paragraph()

    # Disclaimer
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_disc = disc.add_run(
        "Diese Rechnung wurde mit Unterstützung eines KI-Tools erstellt. "
        "Die alleinige Verantwortung für die Richtigkeit und die Einhaltung "
        "des Gerichts- und Notarkostengesetzes (GNotKG) liegt beim Notar.\n"
        f"Fee-Engine: {fee_version}"
    )
    run_disc.font.size = Pt(9)
    run_disc.italic = True

    import io

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _generate_txt(
    positions: list[dict],
    notary: dict,
    now: datetime,
    total_net: float,
    vat: float,
    total_gross: float,
    original_document: str,
    fee_version: str,
) -> bytes:
    lines = []
    lines.append(notary.get("firm_name", ""))
    lines.append(notary.get("name", ""))
    lines.append(notary.get("address", ""))
    lines.append("")
    lines.append("=" * 60)
    lines.append("HONORARRECHNUNG")
    lines.append("=" * 60)
    lines.append(f"Datum: {now.strftime('%d.%m.%Y')}")
    if original_document:
        lines.append(f"Urkunde: {original_document}")
    lines.append("")
    lines.append(f"{'KV-Nr.':10s} {'Beschreibung':30s} {'Wert':>12s} {'Gebühr':>10s}")
    lines.append("-" * 62)
    for p in positions:
        kv = p.get("kv_number", "")
        desc = p.get("description", "")[:28]
        bw = p.get("business_value_eur")
        bw_str = f"{bw:,.2f}" if bw else "–"
        fee_str = f"{p.get('fee_amount', 0.0):,.2f}"
        lines.append(f"{kv:10s} {desc:30s} {bw_str:>12s} {fee_str:>10s} €")
    lines.append("-" * 62)
    lines.append(f"{'':10s} {'':30s} {'Zwischensumme':>12s} {total_net:>10,.2f} €")
    lines.append(f"{'':10s} {'':30s} {'USt 19%':>12s} {vat:>10,.2f} €")
    lines.append(f"{'':10s} {'':30s} {'Gesamtbetrag':>12s} {total_gross:>10,.2f} €")
    lines.append("")
    lines.append(f"Zahlung an: {notary.get('bank_name', '')}")
    lines.append(f"IBAN: {notary.get('iban', '')}")
    if notary.get("bic"):
        lines.append(f"BIC: {notary.get('bic', '')}")
    lines.append("")
    lines.append(
        "Disclaimer: Diese Rechnung wurde mit Unterstützung eines KI-Tools erstellt. "
        "Die alleinige Verantwortung liegt beim Notar."
    )
    lines.append(f"Fee-Engine: {fee_version}")
    return ("\n".join(lines)).encode("utf-8")


def _generate_rtf(
    positions: list[dict],
    notary: dict,
    now: datetime,
    total_net: float,
    vat: float,
    total_gross: float,
    original_document: str,
    fee_version: str,
) -> bytes:
    # RTF als einfaches Text-RTF
    text = _generate_txt(
        positions,
        notary,
        now,
        total_net,
        vat,
        total_gross,
        original_document,
        fee_version,
    ).decode("utf-8")

    rtf = "{\\rtf1\\ansi\\deff0\n"
    rtf += "{\\fonttbl{\\f0\\fswiss Helvetica;}{\\f1\\fmodern Courier;}}\n"
    rtf += "\\f0\\fs24\n"

    for line in text.split("\n"):
        rtf += f"{line}\\par\n"

    rtf += "}"
    return rtf.encode("cp1252", errors="replace")
