"""Dokumenten-Parser: PDF, DOCX, RTF, TXT mit OCR-Fallback."""

from pathlib import Path

from loguru import logger

from core.config import get_settings
from core.models import ExtractionQuality, ParsedDocument


def _parse_txt(file_path: Path) -> ParsedDocument:
    text = file_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    pages = max(1, len(lines) // 60)
    return ParsedDocument(
        full_text=text,
        pages=pages,
        metadata={"format": "txt", "size_bytes": file_path.stat().st_size},
        extraction_quality=ExtractionQuality.GOOD,
    )


def _parse_docx(file_path: Path) -> ParsedDocument:
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    pages = max(1, len(paragraphs) // 40)
    return ParsedDocument(
        full_text=text,
        pages=pages,
        metadata={
            "format": "docx",
            "paragraph_count": len(doc.paragraphs),
            "size_bytes": file_path.stat().st_size,
        },
        extraction_quality=ExtractionQuality.GOOD,
    )


def _parse_rtf(file_path: Path) -> ParsedDocument:
    raw = file_path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("cp1252")
    stripped = _strip_rtf(text)
    stripped = _cleanup_font_residue(stripped)
    lines = stripped.splitlines()
    pages = max(1, len(lines) // 50)
    return ParsedDocument(
        full_text=stripped,
        pages=pages,
        metadata={"format": "rtf", "size_bytes": file_path.stat().st_size},
        extraction_quality=ExtractionQuality.GOOD,
    )


def _strip_rtf(text: str) -> str:
    """Entfernt RTF-Steuerzeichen und liefert Plaintext zurück."""
    import re

    # Hex-Escape \\'XX  →  Latin-1 Zeichen (muss vor anderen Ersetzungen kommen)
    text = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: _rtf_hex_to_char(m.group(1)), text)

    # Unicode-Escape \\uXXXX
    text = re.sub(
        r"\\u(-?\d+)\s*\\?'[0-9a-fA-F]{0,2}",
        lambda m: chr(int(m.group(1))) if 32 <= int(m.group(1)) <= 65535 else "",
        text,
    )

    # Spezielle RTF-Befehle durch Zeichen ersetzen
    # WICHTIG: Doppel-Backslash (\\par = literaler Text) VOR Einfach-Backslash (\par = Befehl)
    text = re.sub(r"\\\\par\b", "\n", text)
    text = re.sub(r"\\par\b", "\n", text)
    text = re.sub(r"\\line\b", "\n", text)
    text = re.sub(r"\\tab\b", "\t", text)
    text = re.sub(r"\\~", "\xa0", text)

    # Alle anderen RTF-Steuerbefehle entfernen
    text = re.sub(r"\\[a-zA-Z]+\d*\s?", "", text)

    # Geschweifte Klammern entfernen
    text = text.replace("{", "").replace("}", "")

    # Semikolons aus fonttbl-Resten entfernen
    text = re.sub(r";+\s*", "", text)

    # Überflüssige Leerzeilen und Leerzeichen entfernen
    text = re.sub(r"\n[ \t]*\n[ \t]*\n+", "\n\n", text)
    text = re.sub(r"^\s+", "", text)
    text = re.sub(r"\s+$", "", text)
    return text


def _remove_rtf_groups(text: str) -> str:
    """Entfernt RTF-Gruppen {} inklusive verschachtelter Gruppen."""
    import re

    result = text
    prev = None
    while prev != result:
        prev = result
        result = re.sub(r"\{[^{}]*\}", "", result)
    return result


def _cleanup_font_residue(text: str) -> str:
    """Entfernt Fontnamen-Resten (Times New Roman, Arial etc.) vom Textanfang."""
    import re

    font_names = [
        r"Times\s*New\s*Roman",
        r"Arial",
        r"Calibri",
        r"Helvetica",
        r"Courier\s*New",
        r"Times",
        r"Tahoma",
        r"Verdana",
        r"Georgia",
        r"Symbol",
        r"Wingdings",
    ]
    for name in font_names:
        text = re.sub(r"^" + name + r"\s*", "", text)

    return text.strip()


def _rtf_hex_to_char(hex_str: str) -> str:
    try:
        return bytes.fromhex(hex_str).decode("latin-1")
    except (ValueError, UnicodeDecodeError):
        return ""


def _rtf_strip_tags(text: str) -> str:
    """Fallback: Aggressivere RTF-Bereinigung."""
    import re

    text = re.sub(r"\\par\b", "\n", text)
    text = re.sub(r"\\\\par\b", "\n", text)
    text = re.sub(r"\\line\b", "\n", text)
    text = re.sub(r"\\tab\b", "\t", text)
    text = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: _rtf_hex_to_char(m.group(1)), text)
    text = re.sub(r"\{\\\*?\\[^}]*\}", "", text)
    text = re.sub(r"\{[^}]*\}", "", text)
    text = re.sub(r"\\[a-zA-Z]+\d*\s?", "", text)
    return text.strip()


def _parse_pdf(file_path: Path) -> ParsedDocument:
    import fitz

    settings = get_settings()
    doc = fitz.open(str(file_path))
    text_parts: list[str] = []
    for page in doc:
        text_parts.append(page.get_text())
    text = "\n".join(text_parts)
    pages = len(doc)
    doc.close()
    parsed = ParsedDocument(
        full_text=text,
        pages=pages,
        metadata={
            "format": "pdf",
            "size_bytes": file_path.stat().st_size,
        },
        extraction_quality=ExtractionQuality.GOOD,
    )

    if len(text.strip()) < 200 and settings.app_ocr_enabled:
        logger.info("Wenig Text im PDF, starte OCR-Fallback …")
        try:
            ocr_text = _ocr_pdf(file_path)
            parsed = ParsedDocument(
                full_text=ocr_text,
                pages=pages,
                metadata={
                    "format": "pdf_ocr",
                    "size_bytes": file_path.stat().st_size,
                },
                extraction_quality=ExtractionQuality.OCR_FALLBACK,
            )
        except Exception as e:
            logger.warning(f"OCR fehlgeschlagen: {e}")
            parsed = ParsedDocument(
                full_text=text,
                pages=pages,
                metadata={
                    "format": "pdf",
                    "size_bytes": file_path.stat().st_size,
                },
                extraction_quality=ExtractionQuality.POOR,
            )
    return parsed


def _ocr_pdf(file_path: Path) -> str:
    import io

    import fitz
    import pytesseract
    from PIL import Image

    settings = get_settings()
    doc = fitz.open(str(file_path))
    text_parts: list[str] = []
    for page in doc:
        pix = page.get_pixmap(dpi=250)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        page_text = pytesseract.image_to_string(img, lang=settings.app_ocr_lang)
        text_parts.append(page_text)
    doc.close()
    return "\n".join(text_parts)


def parse_document(file_path: str | Path) -> ParsedDocument:
    """Hauptfunktion: Dokument parsen und strukturierten Text liefern.

    Args:
        file_path: Pfad zum Dokument (PDF, DOCX, RTF oder TXT).

    Returns:
        ParsedDocument mit Volltext, Seitenzahl und Qualitätsmetadaten.

    Raises:
        ValueError: Bei nicht unterstütztem Format.
        FileNotFoundError: Bei nicht existierender Datei.
    """
    path = Path(file_path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"Datei nicht gefunden: {path}")
    if not path.is_file():
        raise ValueError(f"Pfad ist keine Datei: {path}")

    suffix = path.suffix.lower()
    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".rtf": _parse_rtf,
        ".txt": _parse_txt,
    }

    if suffix not in parsers:
        raise ValueError(
            f"Nicht unterstütztes Format '{suffix}'. Unterstützt: {', '.join(parsers.keys())}"
        )

    file_hash = hash(path.name) % 100000
    logger.info(f"Parse Dokument: #{file_hash} ({suffix})")
    try:
        result = parsers[suffix](path)
        logger.info(
            f"Dokument #{file_hash} geparst: {result.pages} Seiten, "
            f"Qualität: {result.extraction_quality.value}"
        )
        return result
    except Exception as e:
        logger.error(f"Fehler beim Parsen von Dokument #{file_hash}: {type(e).__name__}")
        raise RuntimeError(
            "Fehler beim Parsen des Dokuments. Bitte überprüfen Sie das Dateiformat."
        ) from e
